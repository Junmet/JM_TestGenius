from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional
import logging

from docx import Document
import fitz  # PyMuPDF


logger = logging.getLogger(__name__)
SUPPORTED_SUFFIXES = {".docx", ".md", ".markdown", ".txt", ".pdf"}


@dataclass(frozen=True)
class ParsedDocument:
    path: Path
    text: str


def iter_input_files(input_dir: Path) -> Iterable[Path]:
    """
    遍历 input 目录下所有支持的文件类型，按文件名排序返回 Path 列表。
    """
    for p in sorted(input_dir.glob("*")):
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES:
            logger.debug("已发现支持的输入文件：%s", p)
            yield p


def parse_document(path: Path, *, encoding: str = "utf-8") -> ParsedDocument:
    """
    根据后缀选择不同解析方式：
    - .docx 使用 python-docx 抽取段落和表格；
    - .pdf 使用 PyMuPDF 抽取文本；
    - .md / .txt 直接按文本读取。
    """
    suffix = path.suffix.lower()
    logger.info("正在解析文档：%s（后缀=%s）", path, suffix)
    if suffix == ".docx":
        text = _parse_docx(path)
    elif suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix in {".md", ".markdown", ".txt"}:
        text = _parse_text(path, encoding=encoding)
    else:
        logger.error("不支持的文件类型：%s", path)
        raise ValueError(f"不支持的文件类型：{path.name}")

    text = _normalize_text(text)
    logger.info("文档解析完成：%s，长度=%d 字符", path, len(text))
    return ParsedDocument(path=path, text=text)


def _parse_pdf(path: Path) -> str:
    """
    使用 PyMuPDF 抽取 PDF 文本：逐页提取，保留基本段落换行。
    """
    parts: list[str] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            parts.append(page.get_text())
    finally:
        doc.close()
    text = "\n".join(parts)
    logger.debug("PDF 解析完成：%s，长度=%d 字符", path, len(text))
    return text


def _parse_text(path: Path, *, encoding: str) -> str:
    # 这里不做复杂编码探测：默认 utf-8；如有需要可在 CLI 传入 encoding。
    text = path.read_text(encoding=encoding, errors="ignore")
    logger.debug("文本文件解析完成：%s，长度=%d 字符", path, len(text))
    return text


def _parse_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []

    # 段落：逐段读取，丢弃空行
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    # 表格：将每一行单元格拼成一行文本，避免表格信息丢失
    for table in doc.tables:
        for row in table.rows:
            cells = [(_cell_text(cell.text) or "") for cell in row.cells]
            line = " | ".join([c for c in cells if c.strip()])
            if line.strip():
                parts.append(line.strip())

    text = "\n".join(parts)
    logger.debug("Word（DOCX）解析完成：%s，长度=%d 字符", path, len(text))
    return text


def _cell_text(s: Optional[str]) -> str:
    return (s or "").replace("\xa0", " ").strip()


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 连续空行压缩：把 3 行以上空行压成 2 行，方便 LLM 阅读
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text.strip()

